# /// script
# requires-python = ">=3.10"
# dependencies = ["python-docx", "pypdf", "pyyaml"]
# ///
"""Stage-A form extractor: blank form (.docx | .pdf) -> scheme.yaml field skeleton.

Emits a *skeleton* toward `scheme.yaml` (see references/form-schema-ir.md), not a
finished IR. Every detected control is mapped to a best-guess widget (type-model.md
Axis 1) with any char/word/page limit found nearby. Detections that rest on a guess
are marked `confidence: low` and carry a `note` — this tool locates structure, it does
not resolve meaning (rubric/roles/gates come from the guidelines, done by hand).

    uv run extract_form.py FORM.docx
    uv run extract_form.py FORM.pdf -o skeleton.yaml
"""
from __future__ import annotations

import argparse
import hashlib
import re
import sys
from pathlib import Path

import yaml

W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

# limit patterns -> (unit, regex). Order matters: pages/words/chars.
_LIMIT_PATTERNS = [
    ("pages", re.compile(r"(?:max(?:imum)?|up to|no more than|limit(?:ed)? (?:to|of)?)?\s*(\d{1,3})\s*(?:pages?|pp)\b", re.I)),
    ("words", re.compile(r"(?:max(?:imum)?|up to|no more than|limit(?:ed)? (?:to|of)?)?\s*(\d{2,6})\s*words?\b", re.I)),
    ("chars", re.compile(r"(?:max(?:imum)?|up to|no more than|limit(?:ed)? (?:to|of)?)?\s*(\d{2,6})\s*(?:characters?|chars?)\b", re.I)),
]


def detect_limit(text: str) -> dict | None:
    """Best-guess {value, unit} from surrounding prose. None if nothing plausible."""
    for unit, pat in _LIMIT_PATTERNS:
        m = pat.search(text or "")
        if m:
            return {"value": int(m.group(1)), "unit": unit}
    return None


def _slug(text: str, n: int = 6) -> str:
    # Keep Unicode word chars (don't drop non-ASCII labels to a generic "field"); if a label
    # is punctuation-only / empty, derive a STABLE token from its bytes instead of colliding
    # every such field onto "field".
    words = re.findall(r"\w+", (text or "").lower(), re.UNICODE)
    slug = "-".join(words[:n])
    if not slug:
        return "fld-" + hashlib.sha1((text or "").encode("utf-8")).hexdigest()[:6]
    return slug


class _Unique:
    """Suffix duplicate slugs (-2, -3, …) so every field id in a document is unique."""

    def __init__(self):
        self._seen: dict[str, int] = {}

    def __call__(self, base: str) -> str:
        if base not in self._seen:
            self._seen[base] = 1
            return base
        self._seen[base] += 1
        return f"{base}-{self._seen[base]}"


# ---------------------------------------------------------------- docx ------
def extract_docx(path: Path) -> dict:
    try:
        from docx import Document
        from docx.oxml.ns import qn
    except ImportError:
        sys.exit("error: python-docx not available. Run via `uv run` so PEP 723 deps install.")

    doc = Document(str(path))
    fields: list[dict] = []
    uniq = _Unique()

    # 1. Content controls (SDT). python-docx has no public API for these, so walk
    #    the raw XML. An SDT that is really a form control is high-signal.
    def sdt_type(sdtpr) -> str:
        for tag, widget in (("text", "narrative"), ("checkbox", "boolean-gate"),
                            ("dropDownList", "single-choice"), ("comboBox", "single-choice"),
                            ("date", "scalar"), ("picture", "structured-upload")):
            # w14:checkbox lives in the w14 ns; match by localname to stay ns-agnostic
            if any(el.tag.split("}")[-1] == tag for el in sdtpr.iter()):
                return widget
        return "scalar"

    for sdt in doc.element.body.iter(qn("w:sdt")):
        sdtpr = sdt.find(qn("w:sdtPr"))
        tag = alias = None
        if sdtpr is not None:
            t = sdtpr.find(qn("w:tag"))
            a = sdtpr.find(qn("w:alias"))
            tag = t.get(qn("w:val")) if t is not None else None
            alias = a.get(qn("w:val")) if a is not None else None
        content = "".join(n.text or "" for n in sdt.iter(qn("w:t")))
        label = alias or tag or _slug(content)
        fields.append({
            "id": uniq(_slug(tag or alias or content)),
            "label": label,
            "widget": sdt_type(sdtpr) if sdtpr is not None else "scalar",
            "role": "TODO",
            "limit": detect_limit(content),
            "_source": "content-control",
            "confidence": "medium",
            "note": "content control detected via raw XML; verify widget + role by hand",
        })

    # 2. Mandated headings -> heading-sequenced narrative boundaries. The prose that
    #    follows a heading is the field body; the limit often sits in the heading text.
    #    Skip a heading that only labels a content control already captured above.
    control_labels = {_slug(f["label"]) for f in fields}
    for para in doc.paragraphs:
        style = (para.style.name or "") if para.style else ""
        if style.startswith("Heading") and para.text.strip() and _slug(para.text) not in control_labels:
            fields.append({
                "id": uniq(_slug(para.text)),
                "label": para.text.strip(),
                "widget": "narrative",
                "role": "TODO",
                "limit": detect_limit(para.text),
                "_source": f"heading ({style})",
                "confidence": "low",
                "note": "inferred from a heading; widget=narrative is a default guess",
            })

    # 3. Tables -> likely matrices / repeating groups; never a plain narrative.
    for i, tbl in enumerate(doc.tables):
        header = " | ".join(c.text.strip() for c in tbl.rows[0].cells) if tbl.rows else ""
        fields.append({
            "id": f"table-{i+1}",
            "label": header or f"table {i+1}",
            "widget": "budget-matrix" if re.search(r"\$|cash|year|amount|budget", header, re.I) else "repeating-group",
            "role": "TODO",
            "limit": None,
            "_source": f"table {tbl_dims(tbl)}",
            "confidence": "low",
            "note": "table detected; disambiguate budget-matrix vs repeating-group vs relational-table by hand",
        })

    return {
        "scheme": path.stem,
        "modality": "docx",
        "detected_via": "python-docx (content-controls + headings + tables)",
        "field_count": len(fields),
        "sections": [{"id": "extracted", "title": "Auto-extracted fields", "fields": fields}],
    }


def tbl_dims(tbl) -> str:
    try:
        return f"{len(tbl.rows)}x{len(tbl.columns)}"
    except Exception:
        return "?x?"


# ---------------------------------------------------------------- pdf -------
def extract_pdf(path: Path) -> dict:
    try:
        from pypdf import PdfReader
    except ImportError:
        sys.exit("error: pypdf not available. Run via `uv run` so PEP 723 deps install.")

    reader = PdfReader(str(path))
    modality = "pdf-flat"
    fields: list[dict] = []

    # XFA? (dynamic forms — pypdf can see the packet but cannot reliably fill it)
    try:
        if reader.xfa:
            modality = "pdf-xfa"
    except Exception:
        pass

    acro = None
    try:
        acro = reader.get_fields()
    except Exception:
        acro = None

    if acro:
        modality = "pdf-acroform" if modality != "pdf-xfa" else "pdf-xfa"
        for name, f in acro.items():
            ft = str(f.get("/FT", ""))
            maxlen = f.get("/MaxLen")
            widget, note = _acro_widget(ft, maxlen, f)
            fields.append({
                "id": _slug(str(name)),
                "label": str(f.get("/TU") or name),
                "field_name": str(name),     # exact AcroForm name — render_pdf matches on this
                "widget": widget,
                "role": "TODO",
                "limit": {"value": int(maxlen), "unit": "chars"} if maxlen else None,
                "_source": f"AcroForm {ft}",
                "confidence": "medium",
                "note": note,
            })
    else:
        # No AcroForm: flat text or scanned? Extractable text decides.
        text_len = 0
        try:
            text_len = sum(len(p.extract_text() or "") for p in reader.pages)
        except Exception:
            pass
        per_page = text_len / max(1, len(reader.pages))
        if per_page < 40:  # essentially no selectable text -> image-only
            modality = "pdf-scanned"
        elif modality != "pdf-xfa":
            modality = "pdf-flat"
        fields.append({
            "id": "flat-document",
            "label": path.stem,
            "widget": "structured-upload",
            "role": "TODO",
            "limit": None,
            "_source": f"{len(reader.pages)} pages, ~{per_page:.0f} chars/page",
            "confidence": "low",
            "note": ("no fillable fields; scanned/image PDF — needs OCR + manual field mapping"
                     if modality == "pdf-scanned"
                     else "no AcroForm; flat PDF — fields must be identified from the layout by hand"),
        })

    return {
        "scheme": path.stem,
        "modality": modality,
        "detected_via": "pypdf (AcroForm / XFA / text-density heuristic)",
        "page_count": len(reader.pages),
        "field_count": len([f for f in fields if f["id"] != "flat-document"]),
        "sections": [{"id": "extracted", "title": "Auto-extracted fields", "fields": fields}],
    }


def _acro_widget(ft: str, maxlen, f) -> tuple[str, str]:
    if ft == "/Btn":
        return "boolean-gate", "button field; boolean-gate vs declaration depends on the guidelines"
    if ft == "/Ch":
        ff = int(f.get("/Ff", 0) or 0)
        multi = bool(ff & (1 << 21))  # MultiSelect flag
        return ("multi-choice" if multi else "single-choice"), "choice field; option set is round-bound"
    if ft == "/Sig":
        return "declaration", "signature field"
    if ft == "/Tx":
        if maxlen and int(maxlen) <= 120:
            return "scalar", "short text field; could be scalar or a short narrative"
        return "narrative", "text field; narrative assumed (no/large MaxLen)"
    return "scalar", f"unrecognised field type {ft!r}; defaulted to scalar"


# --------------------------------------------------------------- main -------
def main() -> None:
    ap = argparse.ArgumentParser(description=__doc__, formatter_class=argparse.RawDescriptionHelpFormatter)
    ap.add_argument("form", type=Path, help="blank form: .docx or .pdf")
    ap.add_argument("-o", "--out", type=Path, help="write YAML skeleton here (default: stdout)")
    args = ap.parse_args()

    if not args.form.exists():
        sys.exit(f"error: file not found: {args.form}")

    ext = args.form.suffix.lower()
    if ext == ".docx":
        result = extract_docx(args.form)
    elif ext == ".pdf":
        result = extract_pdf(args.form)
    else:
        sys.exit(f"error: unsupported extension {ext!r} (want .docx or .pdf)")

    header = ("# scheme.yaml SKELETON — auto-extracted, NOT a finished IR.\n"
              "# role: is TODO everywhere; limits/widgets are best-guesses (see confidence/note).\n"
              "# Reconcile against the guidelines by hand before drafting (form-schema-ir.md).\n")
    dumped = header + yaml.safe_dump(result, sort_keys=False, allow_unicode=True)

    low = sum(1 for s in result["sections"] for f in s["fields"] if f.get("confidence") == "low")
    if args.out:
        args.out.write_text(dumped, encoding="utf-8")
        print(f"wrote {args.out}  ({result['field_count']} fields, modality={result['modality']}, {low} low-confidence)")
    else:
        sys.stdout.write(dumped)
        print(f"\n# {result['field_count']} fields, modality={result['modality']}, {low} low-confidence detections", file=sys.stderr)


if __name__ == "__main__":
    main()
