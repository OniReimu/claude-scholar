# /// script
# requires-python = ">=3.10"
# dependencies = ["python-docx", "pypdf", "pyyaml"]
# ///
"""Stage-D pdf renderer: route by PDF sub-type, fill honestly or degrade honestly.

  pdf-acroform             -> VALIDATE fields, then fill with pypdf, save an OFFICIAL PDF.
  pdf-xfa / flat / scanned -> DO NOT fake an official fill. Emit PASTE-READY.txt (C1 grammar)
                              + a companion .docx, both marked NON-OFFICIAL.

Honesty / fail-closed rules:
  * C4 — unresolved IR fields → NON-ZERO exit, DO NOT write an OFFICIAL partial, unless
    --allow-partial (then output is clearly marked NON-OFFICIAL).
  * C4 — before writing an AcroForm: inspect /Ff (required), /MaxLen, /FT (type), /Opt
    (choices); FAIL BEFORE writing on empty-required / over-MaxLen / wrong-type / invalid-choice.
  * C4 — on a fill exception, preserve the true subtype+reason (`pdf-acroform-fill-failed`);
    never relabel as `pdf-flat`.
  * C6 — CONTENT comes from values.yaml (a flat {field-id: value} map), STRUCTURE + limits
    from scheme.yaml (`sections[].fields[]`). A value not in scheme, or a required scheme
    field with no value, is a resolution failure → non-zero (unless --allow-partial).

    uv run render_pdf.py values.yaml form.pdf -o out_dir/ --scheme scheme.yaml
    uv run render_pdf.py filled.yaml form.pdf -o out_dir/          # back-compat flat map

IR/values format: flat mapping | {fields:{...}} | [{id,text}].
"""
from __future__ import annotations

import argparse
import re
import sys
from pathlib import Path

import yaml


def norm(s: str) -> str:
    return re.sub(r"[^a-z0-9]+", "", (s or "").lower())


def load_values(path: Path) -> dict[str, str]:
    raw = yaml.safe_load(path.read_text(encoding="utf-8"))
    if isinstance(raw, dict) and "fields" in raw and isinstance(raw["fields"], dict):
        raw = raw["fields"]
    out: dict[str, str] = {}
    if isinstance(raw, dict):
        for k, v in raw.items():
            out[str(k)] = "" if v is None else str(v)
    elif isinstance(raw, list):
        for item in raw:
            if isinstance(item, dict) and "id" in item:
                out[str(item["id"])] = str(item.get("text", item.get("value", "")))
    else:
        sys.exit("error: values must be a mapping or a list of {id,text} objects")
    return out


def load_scheme(path: Path) -> dict[str, dict]:
    """scheme.yaml (sections[].fields[]) -> {field-id: fieldspec}."""
    data = yaml.safe_load(path.read_text(encoding="utf-8")) or {}
    out: dict[str, dict] = {}

    def walk(fields):
        for f in fields or []:
            fid = f.get("id")
            if fid:
                out[str(fid)] = f
            walk(f.get("sub_fields"))
    for sec in data.get("sections", []):
        walk(sec.get("fields"))
    return out


def scheme_meta(scheme: dict | None, fid: str) -> tuple[str, int | None, str]:
    """(label, limit_value|None, unit) for the C1 header. Default: id, null, chars."""
    if scheme and fid in scheme:
        f = scheme[fid]
        lim = f.get("limit") or {}
        val = lim.get("value") if isinstance(lim, dict) else None
        unit = str(lim.get("unit", "chars")) if isinstance(lim, dict) else "chars"
        return str(f.get("label", fid)), (None if val is None else int(val)), unit
    return fid, None, "chars"


def scheme_required(f: dict) -> bool:
    return bool(f.get("required"))


def resolve_against_scheme(values, scheme):
    """C6 resolution: unknown value ids + required scheme fields with no value."""
    unknown = [k for k in values if k not in scheme]
    missing = [fid for fid, f in scheme.items()
               if scheme_required(f) and not str(values.get(fid, "")).strip()]
    return unknown, missing


def detect_subtype(reader) -> str:
    try:
        if reader.xfa:
            return "pdf-xfa"
    except Exception:
        pass
    try:
        if reader.get_fields():
            return "pdf-acroform"
    except Exception:
        pass
    try:
        chars = sum(len(p.extract_text() or "") for p in reader.pages)
        if chars / max(1, len(reader.pages)) < 40:
            return "pdf-scanned"
    except Exception:
        pass
    return "pdf-flat"


# ---------------------------------------------------------------- acroform ---
def _opt_values(f) -> set[str] | None:
    opt = f.get("/Opt")
    if not opt:
        return None
    vals: set[str] = set()
    for o in opt:
        if isinstance(o, (list, tuple)) and o:
            vals.add(str(o[0]))
            vals.add(str(o[-1]))
        else:
            vals.add(str(o))
    return vals


def validate_acroform(form_fields: dict, mapping: dict[str, str]) -> list[str]:
    """Inspect /Ff //MaxLen //FT //Opt BEFORE writing. Return list of blocking errors."""
    errors: list[str] = []
    # 1. per-value checks over what we are about to write
    for name, val in mapping.items():
        f = form_fields.get(name, {})
        ft = str(f.get("/FT", ""))
        ff = int(f.get("/Ff", 0) or 0)
        required = bool(ff & 2)
        maxlen = f.get("/MaxLen")
        sval = "" if val is None else str(val)
        if required and not sval.strip():
            errors.append(f"empty-required: {name} is required (/Ff Required) but value is blank")
        if maxlen and len(sval) > int(maxlen):
            errors.append(f"over-MaxLen: {name} value {len(sval)} chars > /MaxLen {int(maxlen)}")
        if ft == "/Sig":
            errors.append(f"wrong-type: {name} is a /Sig signature field; cannot fill with text")
        opts = _opt_values(f)
        if ft in ("/Ch", "/Btn") and opts is not None and sval and sval not in opts:
            errors.append(f"invalid-choice: {name}={sval!r} not in /Opt {sorted(opts)}")
    # 2. required form fields we were NOT given a value for
    for name, f in form_fields.items():
        ff = int(f.get("/Ff", 0) or 0)
        if (ff & 2) and name not in mapping:
            errors.append(f"empty-required: {name} is required (/Ff Required) but no value supplied")
    return errors


def resolve_mapping(form_fields, values):
    by_norm = {norm(name): name for name in form_fields}
    mapping, unresolved = {}, []
    for key, val in values.items():
        target = key if key in form_fields else by_norm.get(norm(key))
        if target:
            mapping[target] = val
        else:
            unresolved.append(key)
    return mapping, unresolved


def fill_acroform(reader, mapping: dict[str, str], out_pdf: Path) -> None:
    from pypdf import PdfWriter

    writer = PdfWriter()
    writer.append(reader)
    try:
        writer.set_need_appearances_writer(True)
    except Exception:
        pass
    for page in writer.pages:
        try:
            writer.update_page_form_field_values(page, mapping, auto_regenerate=False)
        except TypeError:
            writer.update_page_form_field_values(page, mapping)
    with out_pdf.open("wb") as fh:
        writer.write(fh)


# ---------------------------------------------------------------- degrade ----
def write_paste_ready(values, out_txt, subtype, scheme, banner_reason) -> None:
    lines = [
        "=" * 78,
        "PASTE-READY — NON-OFFICIAL. Portal / manual entry REQUIRED.",
        f"Source PDF sub-type: {subtype}. {banner_reason}",
        "Values are laid out in the canonical PASTE-READY grammar (charcount.py --text).",
        "=" * 78,
        "",
    ]
    for fid, val in values.items():
        label, lim, unit = scheme_meta(scheme, fid)
        limtxt = "null" if lim is None else str(lim)
        lines.append(f"=== {fid} | {label} | limit: {limtxt} {unit} ===")
        lines.append(val)
        lines.append(f"=== /{fid} ===")
        lines.append("")
    out_txt.write_text("\n".join(lines), encoding="utf-8")


def write_companion_docx(values, out_docx, subtype, scheme) -> bool:
    try:
        from docx import Document
    except ImportError:
        return False
    doc = Document()
    doc.add_heading("NON-OFFICIAL companion — manual/portal entry required", level=0)
    p = doc.add_paragraph()
    p.add_run(f"Source PDF sub-type: {subtype}. ").bold = True
    p.add_run("This .docx is a readable copy of the field values, not a submission artifact.")
    for fid, val in values.items():
        label, lim, unit = scheme_meta(scheme, fid)
        cap = f"{label} ({fid})" + ("" if lim is None else f"  [limit {lim} {unit}]")
        doc.add_heading(cap, level=2)
        doc.add_paragraph(val)
    doc.save(str(out_docx))
    return True


def degrade(values, out_dir, stem, subtype, scheme, reason, exit_code) -> None:
    out_txt = out_dir / f"{stem}-PASTE-READY.txt"
    out_docx = out_dir / f"{stem}-companion.docx"
    write_paste_ready(values, out_txt, subtype, scheme, reason)
    docx_ok = write_companion_docx(values, out_docx, subtype, scheme)
    print("=" * 78)
    print(f"NON-OFFICIAL OUTPUT — sub-type '{subtype}'. Portal / manual entry REQUIRED.")
    print("Nothing here is a submittable PDF.")
    print("=" * 78)
    print(f"reason: {reason}")
    print(f"  wrote {out_txt}")
    print(f"  wrote {out_docx}" if docx_ok else "  (companion .docx skipped: python-docx unavailable)")
    sys.exit(exit_code)


# --------------------------------------------------------------- main --------
def main() -> None:
    ap = argparse.ArgumentParser(description=__doc__, formatter_class=argparse.RawDescriptionHelpFormatter)
    ap.add_argument("ir", type=Path, nargs="?", help="flat IR/values YAML (back-compat; or use --values)")
    ap.add_argument("pdf", type=Path, help="the .pdf form")
    ap.add_argument("-o", "--out", type=Path, required=True, help="output directory (created if missing)")
    ap.add_argument("--scheme", type=Path, help="scheme.yaml (structure + limits) for C6")
    ap.add_argument("--values", type=Path, help="values.yaml (field-id -> value) for C6")
    ap.add_argument("--allow-partial", action="store_true",
                    help="write despite unresolved/invalid fields, clearly marked NON-OFFICIAL")
    args = ap.parse_args()

    values_path = args.values or args.ir
    if values_path is None:
        ap.error("need a values source: positional IR or --values")
    for p in [values_path, args.pdf] + ([args.scheme] if args.scheme else []):
        if not p.exists():
            sys.exit(f"error: file not found: {p}")
    if args.pdf.suffix.lower() != ".pdf":
        sys.exit("error: form must be a .pdf")

    try:
        from pypdf import PdfReader
    except ImportError:
        sys.exit("error: pypdf not available. Run via `uv run` so PEP 723 deps install.")

    values = load_values(values_path)
    if not values:
        sys.exit("error: values contained no fields")
    scheme = load_scheme(args.scheme) if args.scheme else None

    # C6 resolution against scheme.yaml
    scheme_errors: list[str] = []
    if scheme is not None:
        unknown, missing = resolve_against_scheme(values, scheme)
        scheme_errors += [f"unknown-field: {k!r} not in scheme.yaml" for k in unknown]
        scheme_errors += [f"required-missing: scheme field {fid!r} has no value" for fid in missing]

    args.out.mkdir(parents=True, exist_ok=True)
    reader = PdfReader(str(args.pdf))
    subtype = detect_subtype(reader)
    stem = args.pdf.stem

    if subtype == "pdf-acroform":
        form_fields = reader.get_fields() or {}
        mapping, unresolved = resolve_mapping(form_fields, values)
        acro_errors = validate_acroform(form_fields, mapping)
        blockers = ([f"unresolved: {k} (no AcroForm field matched)" for k in unresolved]
                    + scheme_errors + acro_errors)

        if blockers and not args.allow_partial:
            print("REFUSING to write an OFFICIAL partial — validation failed BEFORE writing:", file=sys.stderr)
            for b in blockers:
                print(f"  - {b}", file=sys.stderr)
            print("Fix the values/scheme, or pass --allow-partial for a NON-OFFICIAL draft.", file=sys.stderr)
            sys.exit(2)

        try:
            if blockers and args.allow_partial:
                out_pdf = args.out / f"{stem}-PARTIAL-NON-OFFICIAL.pdf"
            else:
                out_pdf = args.out / f"{stem}-filled.pdf"
            fill_acroform(reader, mapping, out_pdf)
        except Exception as exc:
            # C4: preserve true subtype + reason; DO NOT relabel as pdf-flat
            degrade(values, args.out, stem, "pdf-acroform (fill failed)", scheme,
                    f"pdf-acroform-fill-failed: {exc}", exit_code=4)

        if blockers and args.allow_partial:
            print("=" * 78)
            print(f"NON-OFFICIAL PARTIAL: wrote {out_pdf}")
            print("This is NOT a valid submission — the following did not validate:")
            for b in blockers:
                print(f"  - {b}")
            print("=" * 78)
            sys.exit(3)
        print(f"OFFICIAL fill: wrote {out_pdf}")
        print(f"  filled {len(mapping)}/{len(values)} fields into the AcroForm (validated).")
        return

    # xfa / flat / scanned -> honest degrade
    reason = {
        "pdf-xfa": "XFA dynamic form: pypdf cannot reliably fill the XFA packet.",
        "pdf-flat": "Flat PDF: no form fields to fill.",
        "pdf-scanned": "Scanned/image PDF: no text layer; would need OCR + overlay.",
    }.get(subtype, "Unsupported sub-type.")
    if scheme_errors:
        reason += " (scheme resolution issues: " + "; ".join(scheme_errors) + ")"
    degrade(values, args.out, stem, subtype, scheme, reason, exit_code=3)


if __name__ == "__main__":
    main()
