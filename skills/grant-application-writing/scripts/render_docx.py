# /// script
# requires-python = ">=3.10"
# dependencies = ["python-docx", "pyyaml"]
# ///
"""Stage-D docx renderer: write a filled IR back into the OFFICIAL .docx template.

Fills values into the *same* template so the funder's formatting, headers, and mandated
structure survive. Two write strategies, tried in order per field:
  1. content-control  — match a Word content control (SDT) by TAG then ALIAS; write into it
     with the CORRECT control type (text / checkbox / dropdown / date).
  2. under-heading    — match a mandated heading, insert the value as a paragraph beneath it.

FAIL-CLOSED rules (C5):
  * Type-specific SDT: a checkbox/dropdown/date content-control gets a type-correct write.
    If a value cannot be written as the correct control type (e.g. free text into a checkbox,
    a value not in a dropdown's option list) the field is marked UNRESOLVED and the run exits
    NON-ZERO — never write text into a checkbox and call it done.
  * Multi-paragraph answers are split into paragraphs/runs, preserving the template style.
  * Match by TAG, then ALIAS; a colliding/ambiguous match is reported.
  * A field whose target cannot be located is REPORTED (never silently dropped) → non-zero.

C6: CONTENT from values.yaml (flat {field-id: value}); STRUCTURE from scheme.yaml. A value
not in scheme, or a required scheme field with no value, is a resolution failure → non-zero.

    uv run render_docx.py values.yaml template.docx -o application.docx --scheme scheme.yaml
    uv run render_docx.py filled.yaml template.docx -o application.docx     # back-compat flat map

IR/values format: flat mapping | {fields:{...}} | [{id,text}].
"""
from __future__ import annotations

import argparse
import re
import sys
from pathlib import Path

import yaml

W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
W14 = "http://schemas.microsoft.com/office/word/2010/wordml"
XMLSPACE = "{http://www.w3.org/XML/1998/namespace}space"

_TRUTHY = {"yes", "true", "1", "x", "checked", "on", "y"}
_FALSY = {"no", "false", "0", "unchecked", "off", "n", ""}


def norm(s: str) -> str:
    return re.sub(r"[^a-z0-9]+", "", (s or "").lower())


def load_values(path: Path) -> dict[str, str]:
    raw = yaml.safe_load(path.read_text(encoding="utf-8"))
    if isinstance(raw, dict) and "fields" in raw and isinstance(raw["fields"], dict):
        raw = raw["fields"]
    out: dict[str, str] = {}
    if isinstance(raw, dict):
        for k, v in raw.items():
            out[str(k)] = "" if v is None else str(v)
    elif isinstance(raw, list):
        for item in raw:
            if isinstance(item, dict) and "id" in item:
                out[str(item["id"])] = str(item.get("text", item.get("value", "")))
    else:
        sys.exit("error: values must be a mapping or a list of {id,text} objects")
    return out


def load_scheme(path: Path) -> dict[str, dict]:
    data = yaml.safe_load(path.read_text(encoding="utf-8")) or {}
    out: dict[str, dict] = {}

    def walk(fields):
        for f in fields or []:
            fid = f.get("id")
            if fid:
                out[str(fid)] = f
            walk(f.get("sub_fields"))
    for sec in data.get("sections", []):
        walk(sec.get("fields"))
    return out


def resolve_against_scheme(values, scheme):
    unknown = [k for k in values if k not in scheme]
    missing = [fid for fid, f in scheme.items()
               if bool(f.get("required")) and not str(values.get(fid, "")).strip()]
    return unknown, missing


# ----------------------------------------------------------------- SDT -------
def _attr(el, tagname):
    from docx.oxml.ns import qn
    child = el.find(qn(tagname))
    return child.get(qn("w:val")) if child is not None else None


def _localnames(el):
    return {c.tag.split("}")[-1] for c in el.iter()}


def sdt_kind(sdtpr) -> str:
    lns = _localnames(sdtpr)
    if "checkbox" in lns:
        return "checkbox"
    if "dropDownList" in lns:
        return "dropdown"
    if "comboBox" in lns:
        return "combo"
    if "date" in lns:
        return "date"
    return "text"


def _as_bool(v):
    s = str(v).strip().lower()
    if s in _TRUTHY:
        return True
    if s in _FALSY:
        return False
    return None


def _content_texts(sdt):
    from docx.oxml.ns import qn
    content = sdt.find(qn("w:sdtContent"))
    if content is None:
        return None, []
    return content, list(content.iter(qn("w:t")))


def _set_text(sdt, val) -> bool:
    """Write (possibly multi-paragraph) text; extra lines become <w:br/> in the same run."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    _content, texts = _content_texts(sdt)
    if not texts:
        return False
    lines = str(val).split("\n")
    first = texts[0]
    first.text = lines[0]
    first.set(XMLSPACE, "preserve")
    for extra in texts[1:]:
        extra.text = ""
    run = first.getparent()  # w:r
    for ln in lines[1:]:
        run.append(OxmlElement("w:br"))
        t = OxmlElement("w:t")
        t.text = ln
        t.set(XMLSPACE, "preserve")
        run.append(t)
    return True


def _set_checkbox(sdt, sdtpr, checked: bool) -> None:
    cb = next((c for c in sdtpr.iter() if c.tag.split("}")[-1] == "checkbox"), None)
    if cb is not None:
        chk = next((c for c in cb.iter() if c.tag.split("}")[-1] == "checked"), None)
        if chk is not None:
            chk.set(f"{{{W14}}}val", "1" if checked else "0")
    _content, texts = _content_texts(sdt)
    glyph = "☒" if checked else "☐"   # ☒ / ☐
    if texts:
        texts[0].text = glyph
        for extra in texts[1:]:
            extra.text = ""


def _dropdown_items(sdtpr):
    from docx.oxml.ns import qn
    items = []
    for li in sdtpr.iter(qn("w:listItem")):
        items.append((li.get(qn("w:value")), li.get(qn("w:displayText"))))
    return items


def _set_dropdown(sdt, sdtpr, val) -> bool:
    s = str(val).strip()
    disp = None
    for value, display in _dropdown_items(sdtpr):
        if s in (value, display):
            disp = display or value
            break
    if disp is None:
        return False
    _content, texts = _content_texts(sdt)
    if texts:
        texts[0].text = disp
        for extra in texts[1:]:
            extra.text = ""
    return True


def fill_content_controls(doc, values, filled, unresolved_type, collisions):
    from docx.oxml.ns import qn
    by_norm = {norm(k): k for k in values}
    for sdt in list(doc.element.body.iter(qn("w:sdt"))):
        sdtpr = sdt.find(qn("w:sdtPr"))
        if sdtpr is None:
            continue
        tag, alias = _attr(sdtpr, "w:tag"), _attr(sdtpr, "w:alias")
        key = None
        for ident in (tag, alias):                       # C5: tag first, then alias
            if ident and norm(ident) in by_norm:
                cand = by_norm[norm(ident)]
                if cand in filled:
                    collisions.append((cand, tag, alias))
                else:
                    key = cand
                break
        if not key:
            continue
        kind, val = sdt_kind(sdtpr), values[key]
        if kind == "checkbox":
            b = _as_bool(val)
            if b is None:
                unresolved_type.append((key, "checkbox", val))
                continue
            _set_checkbox(sdt, sdtpr, b)
        elif kind in ("dropdown", "combo"):
            if not _set_dropdown(sdt, sdtpr, val):
                unresolved_type.append((key, kind, val))
                continue
        else:  # text / date  (dates are text-backed)
            if not _set_text(sdt, val):
                unresolved_type.append((key, kind, val))
                continue
        filled.add(key)


def fill_under_headings(doc, values, filled):
    remaining = {norm(k): k for k in values if k not in filled}
    if not remaining:
        return
    for para in doc.paragraphs:
        style = (para.style.name or "") if para.style else ""
        if not style.startswith("Heading"):
            continue
        key = remaining.get(norm(para.text))
        if not key:
            continue
        new_p = para.insert_paragraph_before(str(values[key]).split("\n")[0])
        para._p.addnext(new_p._p)
        try:
            new_p.style = doc.styles["Normal"]
        except KeyError:
            pass
        # additional paragraphs preserve the same style
        anchor = new_p
        for extra in str(values[key]).split("\n")[1:]:
            xp = anchor.insert_paragraph_before(extra)
            anchor._p.addnext(xp._p)
            try:
                xp.style = doc.styles["Normal"]
            except KeyError:
                pass
            anchor = xp
        filled.add(key)
        del remaining[norm(para.text)]


def _clean_paras(body):
    """Split a stored answer into clean paragraphs (blank-line = paragraph; unwrap hard line breaks)."""
    return [re.sub(r"\s+", " ", blk.replace("\n", " ")).strip()
            for blk in str(body).split("\n\n") if blk.strip()]


def fill_under_labels(doc, values, scheme, filled):
    """STRATEGY 3 — tag-less form fill (the AVSTICI/label-based case).

    Many official templates carry NO content controls and NO Heading styles: each field is a plain
    (Normal) LABEL paragraph followed by an empty answer slot. You MUST dissect the template's real
    structure and fill IN PLACE — locate each field's label, then write the answer into the following
    empty paragraph, preserving every label/instruction the template ships. The locator is the scheme
    field's `render_match` (the distinctive label text as it appears in THIS template) or its `label`.
    """
    if not scheme:
        return
    from docx.oxml import OxmlElement
    from docx.text.paragraph import Paragraph
    used_labels: set[str] = set()
    for key in [k for k in values if k not in filled]:
        f = scheme.get(key) or {}
        hint = f.get("render_match") or f.get("label")
        if not hint:
            continue
        h = norm(hint)
        paras = doc.paragraphs                     # re-fetch each field: inserts shift indices
        li = None
        for i, p in enumerate(paras):
            t = norm(p.text)
            if not t or t in used_labels:
                continue
            if t.startswith(h) or (len(h) >= 12 and h in t):
                li = i
                break
        if li is None:
            continue
        ei = next((j for j in range(li + 1, min(li + 7, len(paras)))
                   if paras[j].text.strip() == ""), None)
        if ei is None:
            continue
        parts = _clean_paras(values[key])
        if not parts:
            continue
        paras[ei].text = parts[0]
        anchor = paras[ei]
        for extra in parts[1:]:                    # extra paragraphs INSERTED (never overwrite template text)
            newp = OxmlElement("w:p")
            anchor._p.addnext(newp)
            anchor = Paragraph(newp, anchor._parent)
            anchor.text = extra
        used_labels.add(norm(paras[li].text))
        filled.add(key)


def fill_tables(doc, tables_spec, filled):
    """Fill an official-template TABLE in place (e.g. the budget line-item table). tables_spec:
    {field-id: {header_match?: <substr in the table's header row>, rows: [[cell, ...], ...]}}.
    Rows are added if the template ships fewer; the header row is preserved."""
    for key, spec in (tables_spec or {}).items():
        rows = spec.get("rows") or []
        hm = norm(spec["header_match"]) if spec.get("header_match") else None
        target = None
        for tb in doc.tables:
            if hm is None:
                target = tb
                break
            if tb.rows and hm in norm(" ".join(c.text for c in tb.rows[0].cells)):
                target = tb
                break
        if target is None:
            continue
        while len(target.rows) - 1 < len(rows):
            target.add_row()
        for k, row in enumerate(rows):
            cells = target.rows[k + 1].cells
            for ci, val in enumerate(row[:len(cells)]):
                cells[ci].text = str(val)
        filled.add(key)


def main() -> None:
    ap = argparse.ArgumentParser(description=__doc__, formatter_class=argparse.RawDescriptionHelpFormatter)
    ap.add_argument("ir", type=Path, nargs="?", help="flat IR/values YAML (back-compat; or --values)")
    ap.add_argument("template", type=Path, help="official .docx template to write into")
    ap.add_argument("-o", "--out", type=Path, required=True, help="output .docx path")
    ap.add_argument("--scheme", type=Path, help="scheme.yaml (structure + limits) for C6 + under-label render_match")
    ap.add_argument("--values", type=Path, help="values.yaml (field-id -> value) for C6")
    ap.add_argument("--tables", type=Path, help="tables.yaml {field: {header_match?, rows[][]}} to fill a template table (e.g. budget)")
    args = ap.parse_args()

    values_path = args.values or args.ir
    if values_path is None:
        ap.error("need a values source: positional IR or --values")
    for p in [values_path, args.template] + ([args.scheme] if args.scheme else []):
        if not p.exists():
            sys.exit(f"error: file not found: {p}")
    if args.template.suffix.lower() != ".docx":
        sys.exit("error: template must be a .docx")

    try:
        from docx import Document
    except ImportError:
        sys.exit("error: python-docx not available. Run via `uv run` so PEP 723 deps install.")

    values = load_values(values_path)
    if not values:
        sys.exit("error: values contained no fields")

    scheme: dict = {}
    scheme_errors: list[str] = []
    if args.scheme:
        scheme = load_scheme(args.scheme)
        unknown, missing = resolve_against_scheme(values, scheme)
        scheme_errors += [f"unknown-field: {k!r} not in scheme.yaml" for k in unknown]
        scheme_errors += [f"required-missing: scheme field {fid!r} has no value" for fid in missing]

    tables_spec: dict = {}
    if args.tables:
        if not args.tables.exists():
            sys.exit(f"error: file not found: {args.tables}")
        tables_spec = yaml.safe_load(args.tables.read_text(encoding="utf-8")) or {}

    doc = Document(str(args.template))
    filled: set[str] = set()
    unresolved_type: list = []
    collisions: list = []

    fill_content_controls(doc, values, filled, unresolved_type, collisions)
    n_cc = len(filled)
    fill_under_headings(doc, values, filled)
    n_head = len(filled) - n_cc
    fill_under_labels(doc, values, scheme, filled)          # STRATEGY 3: tag-less label forms
    n_label = len(filled) - n_cc - n_head
    fill_tables(doc, tables_spec, filled)                    # official-template tables (e.g. budget)

    unresolved = [k for k in values if k not in filled and k not in {u[0] for u in unresolved_type}]
    doc.save(str(args.out))

    print(f"wrote {args.out}")
    print(f"  filled: {len(filled)}/{len(values)}  (content-control: {n_cc}, under-heading: {n_head}, under-label: {n_label}, tables: {len(tables_spec)})")
    fail = False
    if collisions:
        print(f"  AMBIGUOUS/COLLIDING ({len(collisions)}) — key already filled by another control:")
        for key, tag, alias in collisions:
            print(f"    - {key} (tag={tag!r} alias={alias!r})")
    if unresolved_type:
        fail = True
        print(f"  TYPE-UNRESOLVED ({len(unresolved_type)}) — value not writable as the control's type:")
        for key, kind, val in unresolved_type:
            print(f"    - {key}: cannot write {val!r} into a {kind} control")
    if unresolved:
        fail = True
        print(f"  UNRESOLVED ({len(unresolved)}) — no content-control tag/alias or matching heading:")
        for k in unresolved:
            print(f"    - {k}")
    if scheme_errors:
        fail = True
        print(f"  SCHEME ERRORS ({len(scheme_errors)}):")
        for e in scheme_errors:
            print(f"    - {e}")
    if fail:
        print("  These did NOT resolve. Fix field-id/type/scheme, or place them by hand.")
        sys.exit(2)


if __name__ == "__main__":
    main()
